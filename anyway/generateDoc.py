from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


def replace_text_in_run(paragraph, run, key, value, style_func):
    if key in run.text:
        parts = run.text.split(key)
        run.text = parts[0]  # 保留占位符前的文本
        if style_func:
            # 创建一个新的 run 并应用样式
            new_run = paragraph.add_run(value)
            style_func(new_run, value)
        else:
            # 直接添加替换的文本
            new_run = paragraph.add_run(value)
        paragraph.add_run(parts[1])  # 添加占位符后的文本


def style_name(run, text):
    run.bold = True
    run.font.name = '仿宋'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    run.font.size = Pt(12)


def style_position(run, text):
    run.italic = True
    run.font.name = 'Arial'
    run.font.size = Pt(14)


def generate_word_from_template(template_path, output_path, data, styles):
    # 加载模板
    doc = Document(template_path)

    # 替换段落中的占位符
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for key, value in data.items():
                if key in run.text:
                    replace_text_in_run(paragraph, run, key, value, styles.get(key))

    # 替换表格中的占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for key, value in data.items():
                            if key in run.text:
                                replace_text_in_run(paragraph, run, key, value, styles.get(key))

    # 保存新的文档
    doc.save(output_path)


# 示例数据
data = {
    '{{name}}': '萧铭聪',
    '{{position}}': '职业陪玩工程师',
    '{{villa}}': '米岗一环中心',
}

# 占位符对应的样式函数
styles = {
    '{{name}}': style_name,
    '{{position}}': style_position,
    '{{villa}}': None,  # 不应用任何样式
}

# 调用函数
template_path = r'C:\Users\Administrator\PycharmProjects\pythonProject1\my_doc\template.docx'
output_path = r'C:\Users\Administrator\PycharmProjects\pythonProject1\my_doc\output.docx'

# 调用函数
generate_word_from_template(template_path, output_path, data, styles)
